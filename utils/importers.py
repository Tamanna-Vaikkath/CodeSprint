"""
Parsers that turn uploaded files of various formats (JSON, TXT, CSV, Excel,
Word, PDF) into a plain list of question dicts, ready to be handed to
utils.storage.add_questions_from_upload().

Each parse_* function returns a tuple: (questions: list[dict], errors: list[str])
`errors` here are *parsing* errors (bad format, unreadable file, etc.) and are
separate from the field-level validation errors produced later by
storage.validate_question().
"""

import csv
import io
import json
import re

TAGS_SPLIT_RE = re.compile(r"[;,]")

# Section markers used by the "rich" Word question-bank format, where each
# question is a Heading-1 paragraph followed by a fixed set of labelled
# sections (Topic, Difficulty Level, Problem Statement, Starter Code, ...).
RICH_DOCX_MARKERS = [
    "Topic:",
    "Sub-Topic:",
    "Difficulty Level:",
    "Problem Statement:",
    "Starter Code (Boilerplate):",
    "Solution Code:",
    "Explanation / Evaluation Notes:",
    "Language:",
    "Question Type:",
    "Points:",
]

# Column headers used by the MCQ Excel template (Interview-Mocha-style).
MCQ_TEMPLATE_HEADERS = {
    "Question Type",
    "Question Text",
    "Option (A)",
    "Option (B)",
    "Correct Answer",
}


def _slugify(text, maxlen=40):
    text = (text or "").lower().strip()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    text = re.sub(r"-+", "-", text).strip("-")
    return text[:maxlen] or "q"


# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------
def _strip_code_fences(text):
    """Remove ```json ... ``` / ``` ... ``` fences some editors add."""
    text = text.strip()
    fence = re.match(r"^```[a-zA-Z]*\s*(.*?)\s*```$", text, re.DOTALL)
    if fence:
        return fence.group(1).strip()
    return text


def _try_parse_json_text(text):
    """Try to interpret raw text as a JSON question bank. Returns list or None."""
    text = _strip_code_fences(text)
    if not text:
        return None
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        return None
    if isinstance(data, dict):
        data = [data]
    if isinstance(data, list):
        return data
    return None


def _row_to_question(row):
    """Convert a flat dict of strings (from CSV/Excel/table row) into a question dict."""
    q = {k.strip(): (v.strip() if isinstance(v, str) else v) for k, v in row.items() if k}

    if "tags" in q and isinstance(q["tags"], str):
        q["tags"] = [t.strip() for t in TAGS_SPLIT_RE.split(q["tags"]) if t.strip()]

    if "test_cases" in q and isinstance(q["test_cases"], str) and q["test_cases"]:
        try:
            q["test_cases"] = json.loads(q["test_cases"])
        except json.JSONDecodeError:
            # leave as-is; validate_question will flag it clearly
            pass

    if "time_limit" in q and isinstance(q["time_limit"], str):
        try:
            q["time_limit"] = int(float(q["time_limit"]))
        except ValueError:
            pass

    return q


# ----------------------------------------------------------------------------
# Per-format parsers
# ----------------------------------------------------------------------------
def parse_json_bytes(raw_bytes):
    text = raw_bytes.decode("utf-8")
    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        return [], [f"Invalid JSON: {e}"]
    if isinstance(data, dict):
        data = [data]
    if not isinstance(data, list):
        return [], ["The JSON file must contain an object or an array of question objects."]
    return data, []


def parse_txt_bytes(raw_bytes):
    text = raw_bytes.decode("utf-8", errors="ignore")

    # 1) maybe it's JSON saved with a .txt extension
    data = _try_parse_json_text(text)
    if data is not None:
        return data, []

    # 2) maybe it's CSV/TSV content saved as .txt
    questions, errors = _parse_delimited_text(text)
    if questions:
        return questions, errors

    return [], [
        "Could not parse the .txt file. It must contain either a JSON array of "
        "questions, or delimited (CSV/TSV-style) rows with a header matching the "
        "question template."
    ]


def _parse_delimited_text(text):
    sample = text[:2048]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
    except csv.Error:
        dialect = csv.excel
    reader = csv.DictReader(io.StringIO(text), dialect=dialect)
    if not reader.fieldnames:
        return [], []
    rows = [_row_to_question(row) for row in reader]
    return rows, []


def parse_csv_bytes(raw_bytes):
    text = raw_bytes.decode("utf-8-sig", errors="ignore")
    questions, errors = _parse_delimited_text(text)
    if not questions:
        errors.append(
            "No rows found in the CSV file. Expected a header row with columns "
            "like id, title, difficulty, tags, description, starter_code, "
            "test_cases, time_limit."
        )
    return questions, errors


def _looks_like_mcq_template(header):
    return MCQ_TEMPLATE_HEADERS.issubset({h for h in header if h})


def _parse_mcq_excel_rows(header, data_rows):
    idx = {h: i for i, h in enumerate(header) if h}

    def cell(row, name):
        i = idx.get(name)
        if i is None or i >= len(row):
            return None
        return row[i]

    questions = []
    for ri, row in enumerate(data_rows):
        qtext = cell(row, "Question Text")
        if qtext is None or not str(qtext).strip():
            continue
        qtext = str(qtext).strip()

        options = []
        for label in ["A", "B", "C", "D", "E"]:
            val = cell(row, f"Option ({label})")
            if val is None:
                continue
            val = str(val).strip()
            if not val or val.lower() == "none":
                continue
            options.append({"label": label, "text": val})

        correct = cell(row, "Correct Answer")
        correct = str(correct).strip().upper() if correct is not None else ""

        difficulty = str(cell(row, "Difficulty Level") or "Medium").strip()
        if difficulty not in {"Easy", "Medium", "Hard"}:
            difficulty = "Medium"

        score = cell(row, "Score")
        try:
            points = int(float(score))
        except (TypeError, ValueError):
            points = 5

        topics = cell(row, "Topics")
        tags = [t.strip() for t in TAGS_SPLIT_RE.split(str(topics)) if t.strip()] if topics else []

        explanation = cell(row, "Answer Explanation")
        explanation = str(explanation).strip() if explanation else ""

        title_match = re.match(r"\s*(Q\d+[\.\):]?\s*[^\n]*)", qtext)
        title = title_match.group(1).strip() if title_match else qtext.splitlines()[0][:80]

        qid = f"mcq-r{ri + 2}-{_slugify(title)}"

        questions.append(
            {
                "id": qid,
                "title": title,
                "difficulty": difficulty,
                "tags": tags,
                "description": qtext,
                "type": "mcq",
                "options": options,
                "correct_answer": correct,
                "explanation": explanation,
                "points": points,
            }
        )
    return questions


def parse_excel_bytes(raw_bytes):
    try:
        import openpyxl
    except ImportError:
        return [], ["openpyxl is required to read Excel files but is not installed."]

    try:
        wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), data_only=True, read_only=True)
    except Exception as e:
        return [], [f"Could not read Excel file: {e}"]

    ws = wb["Questions"] if "Questions" in wb.sheetnames else wb.worksheets[0]
    rows_iter = ws.iter_rows(values_only=True)
    try:
        header = [str(h).strip() if h is not None else "" for h in next(rows_iter)]
    except StopIteration:
        return [], ["The Excel sheet appears to be empty."]

    data_rows = [row for row in rows_iter if row and any(c not in (None, "") for c in row)]
    if not data_rows:
        return [], ["No data rows found in the Excel file."]

    # 1) MCQ question-bank template (Question Type / Question Text / Option (A..E) / Correct Answer).
    if _looks_like_mcq_template(header):
        questions = _parse_mcq_excel_rows(header, data_rows)
        if questions:
            return questions, []
        return [], ["No valid MCQ rows found in the Excel file."]

    # 2) Fallback: generic coding-question schema (id, title, difficulty, ...).
    questions = []
    for row in data_rows:
        row_dict = {header[i]: row[i] for i in range(min(len(header), len(row))) if header[i]}
        row_dict = {k: ("" if v is None else str(v)) for k, v in row_dict.items()}
        questions.append(_row_to_question(row_dict))

    if not questions:
        return [], ["No data rows found in the Excel file."]
    return questions, []


def _looks_like_rich_coding_docx(paragraphs):
    """Heuristic: has Heading-1 questions and most of the expected section labels."""
    has_heading = any(p.style and p.style.name == "Heading 1" for p in paragraphs)
    if not has_heading:
        return False
    texts = [p.text for p in paragraphs]
    found = sum(1 for m in RICH_DOCX_MARKERS if any(t.startswith(m) for t in texts))
    return found >= 6


def _parse_rich_coding_docx(document):
    """
    Parse the structured Word format where each question is:
      Heading 1: "Question N: Title"
      Topic: ...
      Sub-Topic: ...
      Difficulty Level: ...
      Problem Statement:
        ...
      Starter Code (Boilerplate):
        ...
      Solution Code:
        ...
      Explanation / Evaluation Notes:
        ...
      Language: ...
      Question Type: ...
      Points: ...

    The reference "Solution Code" is executed (stdin empty) to auto-derive the
    expected output for a single test case, since the boilerplate/solution
    scripts already print their own result rather than reading from stdin.
    """
    paragraphs = document.paragraphs
    heading_idxs = [i for i, p in enumerate(paragraphs) if p.style and p.style.name == "Heading 1"]
    if not heading_idxs:
        return []
    heading_idxs.append(len(paragraphs))

    def marker_index(block_texts, marker, start=0):
        for j in range(start, len(block_texts)):
            if block_texts[j].startswith(marker):
                return j
        return None

    def section_text(block_texts, start_marker, end_markers):
        si = marker_index(block_texts, start_marker)
        if si is None:
            return ""
        first_line_rest = block_texts[si][len(start_marker):].strip()
        ei = len(block_texts)
        for m in end_markers:
            mi = marker_index(block_texts, m, start=si + 1)
            if mi is not None:
                ei = min(ei, mi)
        body_lines = block_texts[si + 1:ei]
        lines = ([first_line_rest] if first_line_rest else []) + body_lines
        return "\n".join(lines).strip("\n")

    try:
        from .executor import run_code
    except ImportError:
        run_code = None

    questions = []
    for qi in range(len(heading_idxs) - 1):
        start, end = heading_idxs[qi], heading_idxs[qi + 1]
        block_texts = [p.text for p in paragraphs[start:end]]
        heading = block_texts[0]
        title = heading.split(":", 1)[1].strip() if ":" in heading else heading.strip()

        topic = section_text(block_texts, "Topic:", ["Sub-Topic:"]).strip()
        subtopic = section_text(block_texts, "Sub-Topic:", ["Difficulty Level:"]).strip()
        difficulty = section_text(block_texts, "Difficulty Level:", ["Problem Statement:"]).strip()
        problem_statement = section_text(
            block_texts, "Problem Statement:", ["Starter Code (Boilerplate):"]
        )
        starter_code = section_text(
            block_texts, "Starter Code (Boilerplate):", ["Solution Code:"]
        )
        solution_code = section_text(
            block_texts, "Solution Code:", ["Explanation / Evaluation Notes:"]
        )
        explanation = section_text(
            block_texts, "Explanation / Evaluation Notes:", ["Language:"]
        )
        language = section_text(block_texts, "Language:", ["Question Type:"]).strip() or "Python 3"
        points_text = section_text(block_texts, "Points:", []).strip()
        try:
            points = int(float(points_text))
        except ValueError:
            points = 5

        expected_output = ""
        run_warning = None
        if solution_code.strip() and run_code is not None:
            try:
                result = run_code(solution_code, language, "", timeout=10)
                if result["success"]:
                    expected_output = result["stdout"].strip()
                else:
                    run_warning = (result["stderr"] or "solution produced no output").strip()
            except Exception as e:  # noqa: BLE001
                run_warning = str(e)

        qnum_match = re.match(r"Question\s+(\d+)", heading)
        qnum = qnum_match.group(1) if qnum_match else str(qi + 1)
        qid = f"coding-q{qnum}-{_slugify(title)}"

        tags = [t for t in [topic, subtopic] if t]

        q = {
            "id": qid,
            "title": title or f"Question {qnum}",
            "difficulty": difficulty if difficulty in {"Easy", "Medium", "Hard"} else "Medium",
            "tags": tags,
            "description": problem_statement,
            "starter_code": starter_code or "# your code here\n",
            "time_limit": 5,
            "points": points,
            "type": "coding",
            "explanation": explanation,
            "test_cases": [
                {"input": "", "expected_output": expected_output, "hidden": False}
            ],
        }
        if run_warning:
            q["_import_warning"] = (
                f"Auto-generated expected output may be missing/incomplete "
                f"(solution run: {run_warning}); verify the test case manually."
            )
        questions.append(q)
    return questions


def parse_docx_bytes(raw_bytes):
    try:
        import docx
    except ImportError:
        return [], ["python-docx is required to read Word files but is not installed."]

    try:
        document = docx.Document(io.BytesIO(raw_bytes))
    except Exception as e:
        return [], [f"Could not read Word file: {e}"]

    # 1) Rich structured coding-question format (Heading 1 + labelled sections).
    if _looks_like_rich_coding_docx(document.paragraphs):
        questions = _parse_rich_coding_docx(document)
        if questions:
            warn_count = sum(1 for q in questions if q.pop("_import_warning", None))
            errors = (
                [f"{warn_count} question(s) imported with an unverified auto-generated "
                 f"test case — please spot-check them after import."]
                if warn_count
                else []
            )
            return questions, errors

    # 2) Try treating the full document text as JSON (in case someone pasted
    #    a JSON question bank into a Word document).
    full_text = "\n".join(p.text for p in document.paragraphs)
    data = _try_parse_json_text(full_text)
    if data is not None:
        return data, []

    # 3) Try reading a table: first row = headers, remaining rows = questions.
    questions = []
    for table in document.tables:
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        for row in table.rows[1:]:
            cells = [c.text for c in row.cells]
            row_dict = {header[i]: cells[i] for i in range(min(len(header), len(cells))) if header[i]}
            if any(v.strip() for v in row_dict.values()):
                questions.append(_row_to_question(row_dict))

    if questions:
        return questions, []

    return [], [
        "Could not find a JSON question bank, a recognizable structured question "
        "layout, or a table in the Word document. Paste JSON as plain text, use "
        "a table with a header row matching the question template, or use the "
        "structured 'Question N: ...' + labelled-section layout."
    ]


def parse_pdf_bytes(raw_bytes):
    try:
        import pypdf
    except ImportError:
        return [], ["pypdf is required to read PDF files but is not installed."]

    try:
        reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
        full_text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as e:
        return [], [f"Could not read PDF file: {e}"]

    data = _try_parse_json_text(full_text)
    if data is not None:
        return data, []

    # Fall back to delimited-text parsing in case the PDF was a plain table export.
    questions, errors = _parse_delimited_text(full_text)
    if questions:
        return questions, errors

    return [], [
        "Could not extract a question bank from the PDF. PDFs work best when "
        "they contain a JSON question bank as plain text (e.g. printed from a "
        ".json file). For structured data, prefer CSV, Excel, or Word instead."
    ]


# ----------------------------------------------------------------------------
# Dispatcher
# ----------------------------------------------------------------------------
EXTENSION_PARSERS = {
    "json": parse_json_bytes,
    "txt": parse_txt_bytes,
    "csv": parse_csv_bytes,
    "xlsx": parse_excel_bytes,
    "xls": parse_excel_bytes,
    "docx": parse_docx_bytes,
    "pdf": parse_pdf_bytes,
}

SUPPORTED_UPLOAD_TYPES = list(EXTENSION_PARSERS.keys())


def parse_uploaded_file(filename, raw_bytes):
    """
    Dispatch to the right parser based on file extension.
    Returns (questions: list[dict], errors: list[str]).
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    parser = EXTENSION_PARSERS.get(ext)
    if parser is None:
        return [], [f"Unsupported file type: .{ext}"]
    return parser(raw_bytes)
